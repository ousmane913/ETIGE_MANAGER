from pathlib import Path
import html
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / 'DEPLOIEMENT_PRODUCTION.md'
DOCX_PATH = ROOT / 'DEPLOIEMENT_PRODUCTION.docx'
PDF_PATH = ROOT / 'DEPLOIEMENT_PRODUCTION.pdf'


def source_blocks():
    lines = SOURCE.read_text(encoding='utf-8').splitlines()
    blocks = []
    index = 0
    while index < len(lines):
        line = lines[index]
        if line.startswith('```'):
            code = []
            index += 1
            while index < len(lines) and not lines[index].startswith('```'):
                code.append(lines[index])
                index += 1
            blocks.append(('code', '\n'.join(code)))
        elif line.startswith('|'):
            table = []
            while index < len(lines) and lines[index].startswith('|'):
                if not re.match(r'^\|\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$', lines[index]):
                    table.append([cell.strip() for cell in lines[index].strip('|').split('|')])
                index += 1
            blocks.append(('table', table))
            continue
        elif line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            blocks.append(('heading', level, line[level:].strip()))
        elif line.startswith('- '):
            blocks.append(('bullet', line[2:].strip()))
        elif re.match(r'^\d+\.\s+', line):
            blocks.append(('number', re.sub(r'^\d+\.\s+', '', line)))
        elif line.strip():
            blocks.append(('paragraph', line.strip()))
        index += 1
    return blocks


def clean_inline(text):
    text = re.sub(r'\[([^]]+)\]\(([^)]+)\)', r'\1', text)
    return text.replace('**', '').replace('`', '')


def build_docx(blocks):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = document.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal'].font.size = Pt(9)
    code_style = styles.add_style('DeploymentCode', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(8)
    code_style.paragraph_format.left_indent = Inches(0.2)
    code_style.paragraph_format.space_after = Pt(5)

    for block in blocks:
        kind = block[0]
        if kind == 'heading':
            level, text = block[1], clean_inline(block[2])
            document.add_heading(text, level=min(level, 3))
        elif kind == 'code':
            document.add_paragraph(block[1], style='DeploymentCode')
        elif kind == 'bullet':
            document.add_paragraph(clean_inline(block[1]), style='List Bullet')
        elif kind == 'number':
            document.add_paragraph(clean_inline(block[1]), style='List Number')
        elif kind == 'table':
            rows = block[1]
            if not rows:
                continue
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Table Grid'
            for row_index, row in enumerate(rows):
                for col_index, value in enumerate(row):
                    cell = table.cell(row_index, col_index)
                    cell.text = clean_inline(value)
                    if row_index == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
        else:
            document.add_paragraph(clean_inline(block[1]))
    document.save(DOCX_PATH)


def build_pdf(blocks):
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='DeploymentBody', parent=styles['BodyText'], fontName='Helvetica', fontSize=8.5, leading=11, spaceAfter=5))
    styles.add(ParagraphStyle(name='DeploymentHeading', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=15, leading=18, spaceBefore=10, spaceAfter=7, textColor=colors.HexColor('#172033')))
    styles.add(ParagraphStyle(name='DeploymentSubheading', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=11, leading=14, spaceBefore=8, spaceAfter=5, textColor=colors.HexColor('#a65b00')))
    styles.add(ParagraphStyle(name='DeploymentCode', parent=styles['Code'], fontName='Courier', fontSize=6.8, leading=8.2, leftIndent=7, rightIndent=7, backColor=colors.HexColor('#f1f5f9'), borderPadding=5, spaceAfter=6))
    story = []
    for block in blocks:
        kind = block[0]
        if kind == 'heading':
            style = styles['DeploymentHeading'] if block[1] <= 1 else styles['DeploymentSubheading']
            story.append(Paragraph(html.escape(clean_inline(block[2])), style))
        elif kind == 'code':
            story.append(Preformatted(block[1], styles['DeploymentCode']))
        elif kind == 'bullet':
            story.append(Paragraph('• ' + html.escape(clean_inline(block[1])), styles['DeploymentBody']))
        elif kind == 'number':
            story.append(Paragraph(html.escape(clean_inline(block[1])), styles['DeploymentBody']))
        elif kind == 'table':
            rows = [[Paragraph(html.escape(clean_inline(cell)), styles['DeploymentBody']) for cell in row] for row in block[1]]
            table = Table(rows, repeatRows=1, hAlign='LEFT')
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#172033')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#cbd5e1')),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 5),
                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
            ]))
            story.extend([table, Spacer(1, 5)])
        else:
            story.append(Paragraph(html.escape(clean_inline(block[1])), styles['DeploymentBody']))
    document = SimpleDocTemplate(str(PDF_PATH), pagesize=A4, rightMargin=15 * mm, leftMargin=15 * mm, topMargin=14 * mm, bottomMargin=14 * mm, title='Déploiement production ETIGE Manager')
    document.build(story)


if __name__ == '__main__':
    blocks = source_blocks()
    build_docx(blocks)
    build_pdf(blocks)
    print(DOCX_PATH)
    print(PDF_PATH)
